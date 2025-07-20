import pandas as pd
from docx import Document
from docx.shared import Inches


# Créer un DataFrame
df = pd.DataFrame({
    'Nom': ['Alice', 'Bob', 'Charlie'],
    'Âge': [25, 30, 35],
    'Ville': ['Paris', 'Lyon', 'Marseille']
})

class Report():
    def __init__(self, report_name):
        self.document=Document()
        self.report_name= report_name

    def doc_add_heading(self, heading, level=1):
        self.document.add_heading(heading, level=level)
        return self

    def doc_add_paragraph(self, paragraph):
        self.document.add_paragraph(paragraph)
        return self

    def doc_add_image(self, image_path):
        self.document.add_picture(image_path, width=Inches(5.0))
        return self

    def doc_add_table(self, dataframe):
        df=dataframe
        table = self.document.add_table(rows=1, cols=len(df.columns))
        table.style = 'Light Grid'

        # Ajouter les en-têtes
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = str(col)

        # Ajouter les lignes de données
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, item in enumerate(row):
                row_cells[i].text = str(item)

        return self


    def add_combo(self,heading, paragraph="vide", image_path="vide"):
        self.doc_add_heading(heading)
        if not paragraph=="vide":
            self.doc_add_paragraph(paragraph)
        if not image_path=="vide":
            self.doc_add_image(image_path)

    def report_builder(self, path):
        # Sauvegarder le fichier
        self.document.save(f"{path}{self.report_name}.docx")
